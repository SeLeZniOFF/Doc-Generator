import re
from docx import Document
from typing import Dict, Set

PLACEHOLDER_REGEX = re.compile(r"\{[A-Z0-9_]+\}")

def extract_placeholders(doc_path: str) -> Set[str]:
    doc = Document(doc_path)
    found: Set[str] = set()

    def find_in_paragraphs(paragraphs):
        for p in paragraphs:
            text = p.text
            for m in PLACEHOLDER_REGEX.findall(text):
                found.add(m)

    def find_in_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    find_in_paragraphs(cell.paragraphs)

    find_in_paragraphs(doc.paragraphs)
    find_in_tables(doc.tables)

    # headers/footers (минимально)
    for section in doc.sections:
        find_in_paragraphs(section.header.paragraphs)
        find_in_paragraphs(section.footer.paragraphs)

    return found

def replace_placeholders(doc_path: str, mapping: Dict[str, str], on_missing: str = "keep") -> Document:
    """
    Простая замена плейсхолдеров. Возможна потеря тонкого форматирования, если плейсхолдеры разорваны на разные runs.
    """
    doc = Document(doc_path)

    def replace_in_paragraph(p):
        # Склеиваем текст параграфа, заменяем, пересобираем.
        text = p.text
        def repl(m):
            key = m.group(0)
            if key in mapping:
                return mapping[key]
            if on_missing == "error":
                raise KeyError(f"Missing value for placeholder {key}")
            return key  # keep
        new_text = PLACEHOLDER_REGEX.sub(repl, text)
        # очистить runs и задать один run
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)

    def replace_in_table(t):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for t in doc.tables:
        replace_in_table(t)

    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p)

    return doc
